from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import operator
from functools import reduce
from pathlib import Path
import argparse
from PyQt5.QtWidgets import QApplication
import sys
from google_trans_new import google_translator
from filegui import FileWindow



def mergeListAlternative(lst1: list, lst2: list) -> list:
    return list(reduce(operator.add, zip(lst1, lst2)))


def getText(filename: str) -> list:
    doc = Document(filename)
    return [para.text.strip() for para in doc.paragraphs]

def progressBar(iterable, task):
    total = len(iterable)
    def printProgressBar(iteration):
        percent = ("{0:.1f}").format(100 * (iteration / float(total)))
        filledLength = int(50 * iteration // total)
        bar = '█' * filledLength + '-' * (50 - filledLength)
        print(f'\rProgress {task}: |{bar}| {percent}% Complete', end='\r')
    printProgressBar(0)
    for i, item in enumerate(iterable):
        yield item
        printProgressBar(i + 1)
    print()


def start():
    extractedText = getText(input_doc_path)
    print("Text Extracted!")

    filtered_orig_txt = list(filter(None, extractedText))
    translator = google_translator()
    translated_txt = [translator.translate(txt, lang_src=args.src, lang_tgt=args.dest) for txt in progressBar(filtered_orig_txt, 'translate')]
    print('Translated!')
    combined_txt = mergeListAlternative(filtered_orig_txt, translated_txt)
    print('Merged!')

    output_doc = Document()
    styles = output_doc.styles
    styleE = styles.add_style('English', WD_STYLE_TYPE.PARAGRAPH)
    styleE.font.name = 'Times New Roman'
    styleE.font.bold = True
    styleE.font.size = Pt(11)
    styleK = styles.add_style('Korean', WD_STYLE_TYPE.PARAGRAPH)
    styleK.font.name = 'Times New Roman'
    styleK.font.bold = False
    styleK.font.size = Pt(11)
    for i, txt in enumerate(progressBar(combined_txt, 'Export')):
        p = output_doc.add_paragraph(txt)
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing = 1
        p.style = output_doc.styles['English'] if i % 2 else output_doc.styles['Korean']
    output_doc.save(output_doc_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Translate word Document.')
    parser.add_argument('-i', metavar='i', dest='i', type=str, nargs=1, default='',
                        help='import path')
    parser.add_argument('-o', metavar='o', dest='o', type=str, nargs=1, default='',
                        help='import path')
    parser.add_argument('-src', metavar='src', dest='src', type=str, nargs=1, default='ko',
                        help='language to translate from')
    parser.add_argument('-dest', metavar='dest', dest='dest', type=str, nargs=1, default='en',
                        help='language to translate to')
    args = parser.parse_args()
    if args.i:
        input_doc_path = args.i[0]
    else:
        app = QApplication(sys.argv)
        ex = FileWindow()
        input_doc_path = ex.showDialog()
    print(f"set import path to: {input_doc_path}")

    if args.o:
        output_doc_path = args.o[0]
    else:
        output_doc_path = Path(__file__).parent / 'new' / 'translated.docx'
    print(f"set export path to: {output_doc_path}")
    start()
